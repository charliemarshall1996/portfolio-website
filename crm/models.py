"""Data models for all CRM-related records.

Data description for contained models:
- Company: a lead, prospect, client or former-client company.
- Contact: a person working for a lead, prospect, client or former-client company.
- Interaction: an interaction had with a lead, prospect, client or former-client company
and/or contact.
"""

import decimal
from datetime import timedelta
import uuid
from django.core.files.base import ContentFile
from io import BytesIO
from docx import Document
from modelcluster.models import ClusterableModel
from modelcluster.fields import ParentalKey
from wagtail.models import Orderable
from wagtail.admin.panels import FieldPanel, MultiFieldPanel, InlinePanel
from django.db import models
from django.utils import timezone

# Create your models here.


class Setting(models.Model):
    follow_up_days = models.IntegerField(default=3)

    def save(self, *args, **kwargs):
        self.pk = 1
        super().save(*args, **kwargs)


class Company(models.Model):
    name = models.CharField(max_length=255)
    industry = models.CharField(max_length=255, blank=True)
    website = models.URLField(blank=True)
    phone = models.CharField(max_length=20, blank=True)
    email = models.EmailField(blank=True)
    address = models.TextField(blank=True)
    notes = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    panels = [
        FieldPanel("name"),
        FieldPanel("industry"),
        FieldPanel("website"),
        FieldPanel("phone"),
        FieldPanel("email"),
        FieldPanel("address"),
        FieldPanel("notes"),
    ]

    def __str__(self):
        return self.name

    class Meta:
        verbose_name_plural = "Companies"


class Website(Orderable, ClusterableModel):
    contact = ParentalKey("Contact", related_name="websites")
    url = models.URLField()
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    panels = [
        FieldPanel("url")
    ]


class Contact(ClusterableModel):
    SALUTATION_CHOICES = [
        ("mr", "Mr."),
        ("mrs", "Mrs."),
        ("ms", "Ms."),
        ("dr", "Dr."),
    ]
    STATUS_CHOICES = [
        ("le", "Lead"),
        ("pr", "Prospect"),
        ("cl", "Client"),
        ("fc", "Former Client"),
    ]
    salutation = models.CharField(
        max_length=10, choices=SALUTATION_CHOICES, blank=True)
    first_name = models.CharField(max_length=255)
    last_name = models.CharField(max_length=255, blank=True, null=True)
    company = models.ForeignKey(
        Company, on_delete=models.SET_NULL, null=True, blank=True
    )
    position = models.CharField(max_length=255, blank=True)
    email = models.EmailField(blank=True, null=True, unique=True)
    phone = models.CharField(max_length=20, blank=True, null=True)
    mobile = models.CharField(
        max_length=20, blank=True, null=True)
    linkedin = models.URLField(blank=True)
    status = models.CharField(
        max_length=2, choices=STATUS_CHOICES, default=STATUS_CHOICES[0]
    )
    notes = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    panels = [
        MultiFieldPanel(
            [
                FieldPanel("salutation"),
                FieldPanel("first_name"),
                FieldPanel("last_name"),
            ],
            heading="Name",
        ),
        FieldPanel("company"),
        FieldPanel("position"),
        FieldPanel("status"),
        MultiFieldPanel(
            [
                FieldPanel("email"),
                FieldPanel("phone"),
                FieldPanel("mobile"),
                FieldPanel("linkedin"),
                InlinePanel("websites", label="Website"),
            ],
            heading="Contact Information",
        ),
        FieldPanel("notes"),
    ]
    search_term = models.CharField(max_length=250, null=True, blank=True)

    def __str__(self):
        return f"{self.first_name} {self.last_name}"


class Deal(models.Model):
    STAGE_CHOICES = [
        ("prospect", "Prospect"),
        ("qualification", "Qualification"),
        ("proposal", "Proposal"),
        ("negotiation", "Negotiation"),
        ("closed_won", "Closed Won"),
        ("closed_lost", "Closed Lost"),
    ]

    name = models.CharField(max_length=255)
    company = models.ForeignKey(Company, on_delete=models.CASCADE)
    contact = models.ForeignKey(
        Contact, on_delete=models.SET_NULL, null=True, blank=True
    )
    amount = models.DecimalField(max_digits=10, decimal_places=2)
    stage = models.CharField(
        max_length=20, choices=STAGE_CHOICES, default="prospect")
    expected_close_date = models.DateField()
    probability = models.IntegerField(default=0)
    description = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    panels = [
        FieldPanel("name"),
        FieldPanel("company"),
        FieldPanel("contact"),
        MultiFieldPanel(
            [
                FieldPanel("amount"),
                FieldPanel("stage"),
                FieldPanel("expected_close_date"),
                FieldPanel("probability"),
            ],
            heading="Deal Details",
        ),
        FieldPanel("description"),
    ]

    def __str__(self):
        return f"{self.name} - {self.company}"


class Task(models.Model):
    PRIORITY_CHOICES = [
        ("low", "Low"),
        ("medium", "Medium"),
        ("high", "High"),
    ]

    STATUS_CHOICES = [
        ("not_started", "Not Started"),
        ("in_progress", "In Progress"),
        ("completed", "Completed"),
        ("deferred", "Deferred"),
    ]

    title = models.CharField(max_length=255)
    related_to = models.ForeignKey(
        Deal, on_delete=models.CASCADE, null=True, blank=True
    )
    assigned_to = models.ForeignKey(
        Contact, on_delete=models.SET_NULL, null=True, blank=True
    )
    due_date = models.DateField()
    priority = models.CharField(
        max_length=10, choices=PRIORITY_CHOICES, default="medium"
    )
    status = models.CharField(
        max_length=20, choices=STATUS_CHOICES, default="not_started"
    )
    description = models.TextField(blank=True)
    completed = models.BooleanField(default=False)
    completed_date = models.DateField(null=True, blank=True)

    panels = [
        FieldPanel("title"),
        FieldPanel("related_to"),
        FieldPanel("assigned_to"),
        MultiFieldPanel(
            [
                FieldPanel("due_date"),
                FieldPanel("priority"),
                FieldPanel("status"),
                FieldPanel("completed"),
                FieldPanel("completed_date"),
            ],
            heading="Task Details",
        ),
        FieldPanel("description"),
    ]

    def __str__(self):
        return self.title


class Interaction(ClusterableModel):
    """Track all interactions with contacts and companies in the CRM."""

    MEDIUM_CHOICES = [
        ("ph", "Phone"),
        ("em", "Email"),
        ("vc", "Video Call"),
        ("li", "LinkedIn"),
        ("me", "In-Person Meeting"),
        ("ev", "Event"),
        ("no", "Note"),
        ("ot", "Other"),
    ]

    DIRECTION_CHOICES = [
        ("in", "Incoming"),
        ("out", "Outgoing"),
    ]

    contact = models.ForeignKey(
        Contact,
        related_name="interactions",
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        verbose_name="Contact",
    )

    company = models.ForeignKey(
        Company,
        related_name="interactions",
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        verbose_name="Company",
    )

    medium = models.CharField(
        max_length=3,
        choices=MEDIUM_CHOICES,
        default="em",
        verbose_name="Interaction Medium",
    )

    direction = models.CharField(
        max_length=3, choices=DIRECTION_CHOICES, default="out", verbose_name="Direction"
    )

    subject = models.CharField(
        max_length=200, blank=True, null=True, verbose_name="Subject"
    )

    date = models.DateField(default=timezone.now, verbose_name="Date")

    time = models.TimeField(default=timezone.now, verbose_name="Time")

    duration = models.DurationField(
        blank=True, null=True, verbose_name="Duration")

    detail = models.TextField(blank=True, null=True, verbose_name="Details")

    follow_up = models.DateField(
        blank=True, null=True, verbose_name="Follow-up Date")

    follow_up_completed = models.BooleanField(
        default=False, verbose_name="Follow-up Completed?"
    )

    related_deal = models.ForeignKey(
        Deal,
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        verbose_name="Related Deal",
    )

    panels = [
        MultiFieldPanel(
            [
                FieldPanel("contact"),
                FieldPanel("company"),
                FieldPanel("related_deal"),
            ],
            heading="Related Entities",
        ),
        MultiFieldPanel(
            [
                FieldPanel("medium"),
                FieldPanel("direction"),
                FieldPanel("subject"),
            ],
            heading="Interaction Info",
        ),
        MultiFieldPanel(
            [
                FieldPanel("date"),
                FieldPanel("time"),
                FieldPanel("duration"),
            ],
            heading="Timing",
        ),
        FieldPanel("detail"),
        MultiFieldPanel(
            [
                FieldPanel("follow_up"),
                FieldPanel("follow_up_completed"),
            ],
            heading="Follow-up",
        ),
        InlinePanel("attachments", label="Attachments"),
    ]

    def __str__(self):
        contact_name = self.contact.get_full_name() if self.contact else "No Contact"
        company_name = self.company.name if self.company else "No Company"
        return f"{contact_name} | {company_name} | {self.get_medium_display()} on {self.date}"

    def save(self, *args, **kwargs):
        from crm.models import Setting  # Avoid circular import

        # Set default follow-up date if not specified
        if not self.follow_up and not self.follow_up_completed:
            settings = Setting.objects.first()
            follow_up_days = timezone.timedelta(
                days=settings.follow_up_days if settings else 3
            )
            self.follow_up = timezone.now().date() + follow_up_days

        super().save(*args, **kwargs)

    class Meta:
        ordering = ["-date", "-time"]
        verbose_name = "Interaction"
        verbose_name_plural = "Interactions"


class InteractionAttachment(models.Model):
    """Files attached to interactions (emails, call recordings, meeting notes, etc.)."""

    interaction = ParentalKey(
        Interaction, related_name="attachments", on_delete=models.CASCADE
    )

    file = models.FileField(
        upload_to="interaction_attachments/%Y/%m/%d/", verbose_name="File"
    )

    description = models.CharField(
        max_length=255, blank=True, null=True, verbose_name="Description"
    )

    uploaded_at = models.DateTimeField(
        auto_now_add=True, verbose_name="Uploaded At")

    panels = [
        FieldPanel("file"),
        FieldPanel("description"),
    ]

    def __str__(self):
        return self.description or self.file.name

    class Meta:
        verbose_name = "Attachment"
        verbose_name_plural = "Attachments"


class ServiceType(models.Model):
    name = models.CharField(max_length=100)
    description = models.TextField(blank=True)

    panels = [
        FieldPanel("name"),
        FieldPanel("description"),
    ]

    def __str__(self):
        return self.name

    class Meta:
        verbose_name = "Service Type"
        verbose_name_plural = "Service Types"


class PricingOption(models.Model):
    PRICING_MODELS = [
        ("hourly", "Hourly Rate"),
        ("project", "Per Project"),
        ("record", "Per Record"),
        ("package", "Monthly Package"),
    ]

    name = models.CharField(max_length=100)
    service_type = models.ForeignKey(ServiceType, on_delete=models.CASCADE)
    pricing_model = models.CharField(max_length=20, choices=PRICING_MODELS)
    base_rate = models.DecimalField(max_digits=10, decimal_places=2)
    min_hours = models.IntegerField(default=1, blank=True, null=True)
    max_hours = models.IntegerField(blank=True, null=True)
    active = models.BooleanField(default=True)

    panels = [
        FieldPanel("name"),
        FieldPanel("service_type"),
        FieldPanel("pricing_model"),
        FieldPanel("base_rate"),
        MultiFieldPanel(
            [
                FieldPanel("min_hours"),
                FieldPanel("max_hours"),
            ],
            heading="Hourly Constraints",
            classname="collapsible collapsed",
        ),
        FieldPanel("active"),
    ]

    def __str__(self):
        return f"{self.service_type} - {self.name} ({self.get_pricing_model_display()})"


class PackageTier(models.Model):
    pricing_option = models.ForeignKey(
        PricingOption, on_delete=models.CASCADE, related_name="tiers"
    )
    name = models.CharField(max_length=100)
    description = models.TextField(blank=True)
    monthly_rate = models.DecimalField(max_digits=10, decimal_places=2)
    included_hours = models.IntegerField(default=0)
    overage_rate = models.DecimalField(
        max_digits=10, decimal_places=2, blank=True, null=True
    )

    panels = [
        FieldPanel("name"),
        FieldPanel("description"),
        FieldPanel("monthly_rate"),
        FieldPanel("included_hours"),
        FieldPanel("overage_rate"),
    ]

    def __str__(self):
        return f"{self.pricing_option}: {self.name}"


class Service(ClusterableModel):
    STATUS_CHOICES = [
        ("proposed", "Proposed"),
        ("approved", "Approved"),
        ("in_progress", "In Progress"),
        ("completed", "Completed"),
        ("on_hold", "On Hold"),
        ("cancelled", "Cancelled"),
    ]

    name = models.CharField(max_length=200)
    company = models.ForeignKey(Company, on_delete=models.CASCADE)
    contact = models.ForeignKey(
        Contact, on_delete=models.SET_NULL, null=True, blank=True
    )
    service_type = models.ForeignKey(ServiceType, on_delete=models.PROTECT)
    pricing_option = models.ForeignKey(PricingOption, on_delete=models.PROTECT)
    package_tier = models.ForeignKey(
        PackageTier, on_delete=models.SET_NULL, null=True, blank=True
    )
    status = models.CharField(
        max_length=20, choices=STATUS_CHOICES, default="proposed")
    start_date = models.DateField(null=True, blank=True)
    target_date = models.DateField(null=True, blank=True)
    completion_date = models.DateField(null=True, blank=True)
    notes = models.TextField(blank=True)

    panels = [
        FieldPanel("name"),
        MultiFieldPanel(
            [
                FieldPanel("company"),
                FieldPanel("contact"),
            ],
            heading="Client Information",
        ),
        MultiFieldPanel(
            [
                FieldPanel("service_type"),
                FieldPanel("pricing_option"),
                FieldPanel("package_tier"),
            ],
            heading="Service Details",
        ),
        MultiFieldPanel(
            [
                FieldPanel("status"),
                FieldPanel("start_date"),
                FieldPanel("target_date"),
                FieldPanel("completion_date"),
            ],
            heading="Timing",
        ),
        FieldPanel("notes"),
        InlinePanel("time_entries", label="Time Entries"),
        InlinePanel("deliverables", label="Deliverables"),
    ]

    def __str__(self):
        return f"{self.company}: {self.name} ({self.get_status_display()})"

    def get_total_hours(self):
        """Calculate total hours worked on this service"""
        return self.time_entries.aggregate(total=models.Sum("hours"))["total"] or 0

    def get_hourly_billable_amount(self):
        """Calculate billable amount for hourly services"""
        if self.pricing_option.pricing_model != "hourly":
            return decimal.Decimal("0.00")

        total_hours = self.get_total_hours()
        return total_hours * self.pricing_option.base_rate

    def create_invoice(self):
        """Create a draft invoice for this service"""
        if not self.company:
            raise ValueError(
                "Service must have a company to create an invoice")

        # Get the default invoice template
        template = InvoiceTemplate.objects.filter(is_active=True).first()
        if not template:
            raise ValueError("No active invoice template found")

        invoice = Invoice.objects.create(
            template=template,
            service=self,
            company=self.company,
            contact=self.contact,
            status="draft",
        )

        # Add line items based on service type
        if self.pricing_option.pricing_model == "hourly":
            total_hours = self.get_total_hours()
            InvoiceLineItem.objects.create(
                invoice=invoice,
                description=f"{self.service_type.name} - {total_hours} hours at ${self.pricing_option.base_rate}/hour",
                quantity=total_hours,
                unit_price=self.pricing_option.base_rate,
            )
        elif self.pricing_option.pricing_model == "project":
            InvoiceLineItem.objects.create(
                invoice=invoice,
                description=f"{self.service_type.name} - Project Fee",
                quantity=1,
                unit_price=self.pricing_option.base_rate,
            )
        elif self.pricing_option.pricing_model == "record":
            total_records = (
                self.deliverables.aggregate(
                    total=models.Sum("record_count"))["total"]
                or 0
            )
            InvoiceLineItem.objects.create(
                invoice=invoice,
                description=f"{self.service_type.name} - {total_records} records at ${self.pricing_option.base_rate}/record",
                quantity=total_records,
                unit_price=self.pricing_option.base_rate,
            )
        elif self.pricing_option.pricing_model == "package":
            if self.package_tier:
                InvoiceLineItem.objects.create(
                    invoice=invoice,
                    description=f"{self.service_type.name} - {self.package_tier.name} Package",
                    quantity=1,
                    unit_price=self.package_tier.monthly_rate,
                )

                # Add overage if applicable
                total_hours = self.get_total_hours()
                overage = max(0, total_hours -
                              self.package_tier.included_hours)
                if overage > 0 and self.package_tier.overage_rate:
                    InvoiceLineItem.objects.create(
                        invoice=invoice,
                        description=f"Overage - {overage} additional hours at ${self.package_tier.overage_rate}/hour",
                        quantity=overage,
                        unit_price=self.package_tier.overage_rate,
                    )

        invoice.calculate_totals()
        return invoice


class TimeEntry(models.Model):
    service = ParentalKey(
        Service, on_delete=models.CASCADE, related_name="time_entries"
    )
    date = models.DateField()
    hours = models.DecimalField(max_digits=5, decimal_places=2)
    description = models.TextField()
    billable = models.BooleanField(default=True)

    panels = [
        FieldPanel("date"),
        FieldPanel("hours"),
        FieldPanel("description"),
        FieldPanel("billable"),
    ]

    def __str__(self):
        return f"{self.date}: {self.hours} hours"


class Deliverable(models.Model):
    service = ParentalKey(
        Service, on_delete=models.CASCADE, related_name="deliverables"
    )
    name = models.CharField(max_length=200)
    description = models.TextField(blank=True)
    record_count = models.IntegerField(default=0, blank=True)
    completed = models.BooleanField(default=False)
    completion_date = models.DateField(null=True, blank=True)

    panels = [
        FieldPanel("name"),
        FieldPanel("description"),
        FieldPanel("record_count"),
        FieldPanel("completed"),
        FieldPanel("completion_date"),
    ]

    def __str__(self):
        return self.name


class ContractTemplate(models.Model):
    """Template for generating contracts in Word format"""

    name = models.CharField(max_length=255)
    description = models.TextField(blank=True)
    template_file = models.FileField(upload_to="contract_templates/")
    is_active = models.BooleanField(default=True)

    panels = [
        FieldPanel("name"),
        FieldPanel("description"),
        FieldPanel("template_file"),
        FieldPanel("is_active"),
    ]

    def __str__(self):
        return self.name


class InvoiceTemplate(models.Model):
    """Template for generating invoices in Word format"""

    name = models.CharField(max_length=255)
    description = models.TextField(blank=True)
    template_file = models.FileField(upload_to="invoice_templates/")
    is_active = models.BooleanField(default=True)

    panels = [
        FieldPanel("name"),
        FieldPanel("description"),
        FieldPanel("template_file"),
        FieldPanel("is_active"),
    ]

    def __str__(self):
        return self.name


class Contract(ClusterableModel):
    """Generated contract document"""

    STATUS_CHOICES = [
        ("draft", "Draft"),
        ("sent", "Sent"),
        ("approved", "Approved"),
        ("rejected", "Rejected"),
        ("expired", "Expired"),
    ]

    contract_id = models.CharField(max_length=20, unique=True, editable=False)
    template = models.ForeignKey(ContractTemplate, on_delete=models.PROTECT)
    service = models.ForeignKey(Service, on_delete=models.PROTECT)
    company = models.ForeignKey(Company, on_delete=models.PROTECT)
    contact = models.ForeignKey(
        Contact, on_delete=models.PROTECT, null=True, blank=True
    )
    status = models.CharField(
        max_length=20, choices=STATUS_CHOICES, default="draft")
    generated_date = models.DateField(null=True, blank=True)
    sent_date = models.DateField(null=True, blank=True)
    approved_date = models.DateField(null=True, blank=True)
    expiration_date = models.DateField(null=True, blank=True)
    notes = models.TextField(blank=True)
    document = models.FileField(upload_to="contracts/%Y/%m/", blank=True)

    panels = [
        FieldPanel("template"),
        FieldPanel("service"),
        MultiFieldPanel(
            [
                FieldPanel("company"),
                FieldPanel("contact"),
            ],
            heading="Parties",
        ),
        FieldPanel("status"),
        MultiFieldPanel(
            [
                FieldPanel("generated_date"),
                FieldPanel("sent_date"),
                FieldPanel("approved_date"),
                FieldPanel("expiration_date"),
            ],
            heading="Dates",
        ),
        FieldPanel("notes"),
        FieldPanel("document"),
    ]

    def save(self, *args, **kwargs):
        if not self.contract_id:
            self.contract_id = f"CTR-{uuid.uuid4().hex[:8].upper()}"

        if not self.expiration_date and self.generated_date:
            self.expiration_date = self.generated_date + timedelta(days=30)

        super().save(*args, **kwargs)

    def generate_document(self):
        """Generate Word document from template."""
        doc = Document(self.template.template_file.path)

        # Replace placeholders in the document
        replacements = {
            "{{contract_id}}": self.contract_id,
            "{{date}}": self.generated_date.strftime("%B %d, %Y"),
            "{{company_name}}": self.company.name,
            "{{contact_name}}": self.contact.get_full_name() if self.contact else "",
            "{{service_name}}": self.service.name,
            "{{service_type}}": self.service.service_type.name,
            "{{pricing_model}}": self.service.pricing_option.get_pricing_model_display(),
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        # Save the document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"contract_{self.contract_id}.docx"
        self.document.save(filename, ContentFile(buffer.read()))
        self.save()

    def __str__(self):
        return f"{self.contract_id} - {self.service}"


class Invoice(ClusterableModel):
    """Generated invoice document."""

    STATUS_CHOICES = [
        ("draft", "Draft"),
        ("sent", "Sent"),
        ("paid", "Paid"),
        ("overdue", "Overdue"),
        ("cancelled", "Cancelled"),
    ]

    invoice_id = models.CharField(max_length=20, unique=True, editable=False)
    template = models.ForeignKey(InvoiceTemplate, on_delete=models.PROTECT)
    service = models.ForeignKey(Service, on_delete=models.PROTECT)
    company = models.ForeignKey(Company, on_delete=models.PROTECT)
    contact = models.ForeignKey(
        Contact, on_delete=models.PROTECT, null=True, blank=True
    )
    status = models.CharField(
        max_length=20, choices=STATUS_CHOICES, default="draft")
    issue_date = models.DateField(null=True, blank=True)
    due_date = models.DateField()
    sent_date = models.DateField(null=True, blank=True)
    paid_date = models.DateField(null=True, blank=True)
    notes = models.TextField(blank=True)
    document = models.FileField(upload_to="invoices/%Y/%m/", blank=True)
    subtotal = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    tax = models.DecimalField(max_digits=10, decimal_places=2, default=0)
    total = models.DecimalField(max_digits=10, decimal_places=2, default=0)

    panels = [
        FieldPanel("template"),
        FieldPanel("service"),
        MultiFieldPanel(
            [
                FieldPanel("company"),
                FieldPanel("contact"),
            ],
            heading="Billing Information",
        ),
        FieldPanel("status"),
        MultiFieldPanel(
            [
                FieldPanel("issue_date"),
                FieldPanel("due_date"),
                FieldPanel("sent_date"),
                FieldPanel("paid_date"),
            ],
            heading="Dates",
        ),
        MultiFieldPanel(
            [
                FieldPanel("subtotal"),
                FieldPanel("tax"),
                FieldPanel("total"),
            ],
            heading="Amounts",
        ),
        FieldPanel("notes"),
        FieldPanel("document"),
        InlinePanel("line_items", label="Line Items"),
    ]

    def save(self, *args, **kwargs):
        if not self.invoice_id:
            self.invoice_id = f"INV-{uuid.uuid4().hex[:8].upper()}"

        if not self.due_date and self.issue_date:
            self.due_date = self.issue_date + timedelta(days=14)

        super().save(*args, **kwargs)

    def calculate_totals(self):
        """Calculate subtotal, tax, and total based on line items."""
        self.subtotal = sum(item.amount for item in self.line_items.all())
        # Example tax calculation - adjust as needed
        self.tax = self.subtotal * decimal.Decimal("0.1")  # 10% tax
        self.total = self.subtotal + self.tax
        self.save()

    def generate_document(self):
        """Generate Word document from template."""
        doc = Document(self.template.template_file.path)

        # Calculate totals if not already set
        if self.total == 0:
            self.calculate_totals()

        # Prepare line items table
        line_items = []
        for item in self.line_items.all():
            line_items.append(
                {
                    "description": item.description,
                    "quantity": item.quantity,
                    "unit_price": item.unit_price,
                    "amount": item.amount,
                }
            )

        # Replace placeholders in the document
        replacements = {
            "{{invoice_id}}": self.invoice_id,
            "{{issue_date}}": self.issue_date.strftime("%B %d, %Y"),
            "{{due_date}}": self.due_date.strftime("%B %d, %Y"),
            "{{company_name}}": self.company.name,
            "{{contact_name}}": self.contact.get_full_name() if self.contact else "",
            "{{service_name}}": self.service.name,
            "{{subtotal}}": f"${self.subtotal:,.2f}",
            "{{tax}}": f"${self.tax:,.2f}",
            "{{total}}": f"${self.total:,.2f}",
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        # Find and populate the line items table
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{line_items}}" in cell.text:
                        # Clear the cell
                        cell.text = ""
                        # Add line items
                        for item in line_items:
                            row = table.add_row()
                            row.cells[0].text = item["description"]
                            row.cells[1].text = str(item["quantity"])
                            row.cells[2].text = f"${item['unit_price']:,.2f}"
                            row.cells[3].text = f"${item['amount']:,.2f}"

        # Save the document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"invoice_{self.invoice_id}.docx"
        self.document.save(filename, ContentFile(buffer.read()))
        self.save()

    def __str__(self):
        return f"{self.invoice_id} - {self.service}"


class InvoiceLineItem(models.Model):
    """Line items for an invoice."""

    invoice = ParentalKey(Invoice, on_delete=models.CASCADE,
                          related_name="line_items")
    description = models.CharField(max_length=255)
    quantity = models.DecimalField(max_digits=10, decimal_places=2, default=1)
    unit_price = models.DecimalField(max_digits=10, decimal_places=2)
    amount = models.DecimalField(
        max_digits=10, decimal_places=2, editable=False)

    panels = [
        FieldPanel("description"),
        FieldPanel("quantity"),
        FieldPanel("unit_price"),
    ]

    def save(self, *args, **kwargs):
        self.amount = self.quantity * self.unit_price
        super().save(*args, **kwargs)
        self.invoice.calculate_totals()

    def __str__(self):
        return f"{self.description} - ${self.amount}"
